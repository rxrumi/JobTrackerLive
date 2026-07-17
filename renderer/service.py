import base64
import json
import os
import re
import shutil
import subprocess
import tempfile
from difflib import SequenceMatcher
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader


MAX_REQUEST_BYTES = 2 * 1024 * 1024


def normalize_text(value):
    return re.sub(r"\s+", " ", str(value or "")).strip().lower()


def add_bottom_border(paragraph, color="1f4ed8"):
    props = paragraph._p.get_or_add_pPr()
    borders = props.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        props.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def configure_document(document, template, page_target):
    section = document.sections[0]
    margin = 0.55 if template == "compact" else 0.65
    section.top_margin = Inches(margin)
    section.bottom_margin = Inches(margin)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    normal = document.styles["Normal"]
    normal.font.name = "Liberation Sans"
    normal.font.size = Pt(9.2 if template == "compact" else 10)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.0
    for style_name in ["Resume Section", "Resume Role"]:
        if style_name not in document.styles:
            document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    section_style = document.styles["Resume Section"]
    section_style.font.name = "Liberation Sans"
    section_style.font.size = Pt(10.5)
    section_style.font.bold = True
    section_style.font.all_caps = True
    section_style.paragraph_format.space_before = Pt(6)
    section_style.paragraph_format.space_after = Pt(2)
    role_style = document.styles["Resume Role"]
    role_style.font.name = "Liberation Sans"
    role_style.font.size = Pt(10)
    role_style.font.bold = True
    role_style.paragraph_format.space_before = Pt(4)
    role_style.paragraph_format.space_after = Pt(0)
    return page_target


def add_section_heading(document, title, template):
    paragraph = document.add_paragraph(title.upper(), style="Resume Section")
    if template == "modern":
        paragraph.runs[0].font.color.rgb = RGBColor(31, 78, 216)
    add_bottom_border(paragraph, "1f4ed8" if template == "modern" else "444444")


def add_claim_bullets(document, claims):
    for claim in claims or []:
        text = str(claim.get("text") or "").strip()
        if not text:
            continue
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.18)
        paragraph.paragraph_format.first_line_indent = Inches(-0.12)
        paragraph.paragraph_format.space_after = Pt(1.5)
        paragraph.add_run(text)


def build_docx(canonical, template, page_target, output_path):
    document = Document()
    configure_document(document, template, page_target)
    contact = canonical.get("contact") or {}
    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(1)
    run = name.add_run(contact.get("name") or "Candidate")
    run.bold = True
    run.font.name = "Liberation Sans"
    run.font.size = Pt(17 if template != "compact" else 15)
    if template == "modern":
        run.font.color.rgb = RGBColor(31, 78, 216)
    details = [contact.get(key) for key in ["email", "phone", "location", "linkedin_url"] if contact.get(key)]
    if details:
        line = document.add_paragraph(" | ".join(details))
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line.paragraph_format.space_after = Pt(3)
    if canonical.get("headline"):
        headline = document.add_paragraph(canonical["headline"])
        headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        headline.runs[0].bold = True

    if canonical.get("summary_claims"):
        add_section_heading(document, "Summary", template)
        document.add_paragraph(" ".join(claim.get("text", "") for claim in canonical["summary_claims"] if claim.get("text")))
    if canonical.get("skills"):
        add_section_heading(document, "Skills", template)
        document.add_paragraph(" • ".join(canonical["skills"]))
    if canonical.get("experience"):
        add_section_heading(document, "Experience", template)
        for role in canonical["experience"]:
            heading = " — ".join(part for part in [role.get("title"), role.get("employer")] if part)
            document.add_paragraph(heading, style="Resume Role")
            dates = " – ".join(part for part in [role.get("start_date"), role.get("end_date")] if part)
            meta = " | ".join(part for part in [dates, role.get("location")] if part)
            if meta:
                paragraph = document.add_paragraph(meta)
                paragraph.runs[0].italic = True
            add_claim_bullets(document, role.get("bullets"))
    if canonical.get("projects"):
        add_section_heading(document, "Projects", template)
        for project in canonical["projects"]:
            document.add_paragraph(project.get("name") or "Project", style="Resume Role")
            add_claim_bullets(document, project.get("bullets"))
    if canonical.get("education"):
        add_section_heading(document, "Education", template)
        for item in canonical["education"]:
            document.add_paragraph(" — ".join(part for part in [item.get("credential"), item.get("institution"), item.get("date")] if part))
    if canonical.get("certifications"):
        add_section_heading(document, "Certifications", template)
        for item in canonical["certifications"]:
            document.add_paragraph(" — ".join(part for part in [item.get("name"), item.get("issuer"), item.get("date")] if part))
    document.save(output_path)


def document_text(path):
    document = Document(path)
    return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())


def canonical_claims(canonical):
    claims = [claim.get("text", "") for claim in canonical.get("summary_claims", [])]
    for role in canonical.get("experience", []):
        claims.extend(claim.get("text", "") for claim in role.get("bullets", []))
    for project in canonical.get("projects", []):
        claims.extend(claim.get("text", "") for claim in project.get("bullets", []))
    return [normalize_text(claim) for claim in claims if normalize_text(claim)]


def canonical_display_strings(canonical):
    values = []
    contact = canonical.get("contact") or {}
    values.extend(contact.get(key, "") for key in ["name", "email", "phone", "location", "linkedin_url"])
    values.append(canonical.get("headline", ""))
    values.extend(canonical.get("skills", []))
    values.extend(canonical_claims(canonical))
    for role in canonical.get("experience", []):
        values.extend(role.get(key, "") for key in ["employer", "title", "location", "start_date", "end_date"])
    for project in canonical.get("projects", []):
        values.append(project.get("name", ""))
    for item in canonical.get("education", []):
        values.extend(item.get(key, "") for key in ["institution", "credential", "date"])
    for item in canonical.get("certifications", []):
        values.extend(item.get(key, "") for key in ["name", "issuer", "date"])
    return [normalize_text(value) for value in values if normalize_text(value)]


def render(payload):
    canonical = payload.get("canonical_resume") or {}
    template = payload.get("template") or "classic"
    page_target = int(payload.get("page_target") or 1)
    if template not in {"classic", "compact", "modern"} or page_target not in {1, 2}:
        raise ValueError("invalid render settings")
    with tempfile.TemporaryDirectory(prefix="resume-render-") as directory:
        directory = Path(directory)
        docx_path = directory / "resume.docx"
        pdf_path = directory / "resume.pdf"
        build_docx(canonical, template, page_target, docx_path)
        office_binary = os.environ.get("LIBREOFFICE_BIN") or shutil.which("libreoffice") or shutil.which("soffice")
        if not office_binary:
            raise RuntimeError("LibreOffice is unavailable")
        subprocess.run([
            office_binary, "--headless", "--convert-to", "pdf", "--outdir", str(directory), str(docx_path)
        ], check=True, timeout=45, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        if not pdf_path.exists():
            raise RuntimeError("pdf conversion failed")
        docx_text = document_text(docx_path)
        reader = PdfReader(str(pdf_path))
        pypdf_text = "\n".join(page.extract_text() or "" for page in reader.pages)
        poppler_text = subprocess.run(
            ["pdftotext", str(pdf_path), "-"], check=True, timeout=20, stdout=subprocess.PIPE, stderr=subprocess.PIPE
        ).stdout.decode("utf-8", errors="replace")
        preview_prefix = directory / "preview"
        subprocess.run(
            ["pdftoppm", "-f", "1", "-singlefile", "-png", "-r", "110", str(pdf_path), str(preview_prefix)],
            check=True, timeout=30, stdout=subprocess.PIPE, stderr=subprocess.PIPE
        )
        preview_rendered = (directory / "preview.png").exists()
        pdf_text = poppler_text or pypdf_text
        normalized_docx = normalize_text(docx_text)
        normalized_pdf = normalize_text(pdf_text)
        agreement = SequenceMatcher(None, normalized_docx, normalized_pdf).ratio() if normalized_docx else 0
        claims_present = all(claim in normalized_docx for claim in canonical_claims(canonical))
        canonical_text_present = all(value in normalized_docx for value in canonical_display_strings(canonical))
        page_count = len(reader.pages)
        qa = {
            "passed": bool(normalized_pdf) and agreement >= 0.96 and claims_present and canonical_text_present and page_count <= page_target and preview_rendered,
            "page_count": page_count,
            "page_target": page_target,
            "selectable_pdf_text": bool(normalized_pdf),
            "text_agreement": agreement >= 0.96,
            "text_agreement_ratio": round(agreement, 4),
            "canonical_claims_present": claims_present,
            "canonical_text_present": canonical_text_present,
            "poppler_rendered": preview_rendered,
            "overflow": page_count > page_target,
        }
        return {
            "docx_base64": base64.b64encode(docx_path.read_bytes()).decode("ascii"),
            "pdf_base64": base64.b64encode(pdf_path.read_bytes()).decode("ascii"),
            "qa": qa,
        }


class Handler(BaseHTTPRequestHandler):
    def do_GET(self):
        if self.path == "/health":
            self.send_json(200, {"ok": True})
        else:
            self.send_json(404, {"error": "not_found"})

    def do_POST(self):
        if self.path != "/render":
            self.send_json(404, {"error": "not_found"})
            return
        try:
            length = int(self.headers.get("content-length", "0"))
            if length <= 0 or length > MAX_REQUEST_BYTES:
                self.send_json(413, {"error": "invalid_request_size"})
                return
            payload = json.loads(self.rfile.read(length))
            self.send_json(200, render(payload))
        except (ValueError, json.JSONDecodeError) as error:
            self.send_json(400, {"error": str(error)})
        except Exception:
            self.send_json(500, {"error": "render_failed"})

    def send_json(self, status, payload):
        body = json.dumps(payload, separators=(",", ":")).encode("utf-8")
        self.send_response(status)
        self.send_header("content-type", "application/json")
        self.send_header("content-length", str(len(body)))
        self.send_header("cache-control", "no-store")
        self.end_headers()
        self.wfile.write(body)

    def log_message(self, format_string, *args):
        # Request bodies contain private resume content and are never logged.
        return


if __name__ == "__main__":
    port = int(os.environ.get("PORT", "8080"))
    ThreadingHTTPServer(("0.0.0.0", port), Handler).serve_forever()
